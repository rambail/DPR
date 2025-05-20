"""
    Generates a Word report for the metro corridor.

    Parameters:
        corridor (str): Name of the corridor.
        years (list): List of year labels.
        data (dict): Dictionary with keys 'DailyRidership', 'PHPDT' and corresponding yearly values.
        train (dict): Dictionary with train parameters like DMC, TC, MC, TrainComp.
        parameters (dict): Dictionary with 'AverageSpeed', 'SectionLength', 'ReversalTime'.
        headways (dict): Dictionary of headways per year.
        num_trains (dict): Dictionary of number of trains per year.
        train_composition (str): Composition of train (e.g., "DMC,TC,DMC").
        capacity (int): Train carrying capacity (AW3).
        section (str): Section Header for the report
        tpower (dict): Parameters for calculating Traction power requirements
        apower (dict): Parameters for calculating Auxiliary power requirements

    Returns:
        str: Path to the saved Word document.
    """

import os
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class MetroReportGenerator:
    def __init__(self,
                corridor: str, parameters: dict[str, str|float|int], train_composition: str,
                capacity: dict[str, int|float], years: list[int], tfc_labels: list[str],
                data: dict, yearly_headways: dict[int, float],
                yearly_trains: dict[int, int], section: str,
                tpower: dict[int, float], apower: dict[int, float], pw_labels: list[str],
                trc_energy: dict[int, float], aux_energy: dict[int, float],
                total_energy: dict[int, float],
                image_filename: str, output_dir_dpr: str
            ):
        self.corridor = corridor
        self.parameters = parameters
        self.train_composition = train_composition
        self.capacity = capacity
        self.years = years
        self.tfc_labels = tfc_labels
        self.data = data
        self.yearly_headways = yearly_headways
        self.yearly_trains = yearly_trains
        self.section = section
        self.tpower = tpower
        self.apower = apower
        self.pw_labels = pw_labels
        self.trc_energy = trc_energy
        self.aux_energy = aux_energy
        self.total_energy = total_energy
        self.image_filename = image_filename
        self.output_dir = output_dir_dpr
        self.doc_path = os.path.join(self.output_dir, "Metro_Traffic_Report.docx")

    def set_cell_background(self, cell, color):
        """Set background color of a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def add_formatted_section(self, label, spaces):
        section_header = self.doc.add_heading(label, level=1)
        for run in section_header.runs:
            run.font.underline = True
        spacing = self.doc.add_paragraph()
        spacing.paragraph_format.space_after = Pt(spaces)

    def add_formatted_para(self, label, value, unit=""):
        para = self.doc.add_paragraph()
        para.add_run(label)
        value_run = para.add_run(str(value))
        value_run.bold = True
        if unit:
            para.add_run(f" {unit}")
        para.paragraph_format.line_spacing = 1.5

    def create_formatted_table(self, header, labels, data):
        table = self.doc.add_table(rows=len(labels) + 1, cols=len(header) + 1)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        for i, h in enumerate(header):
            hdr_cells[i + 1].text = str(h)

        for cell in hdr_cells:
            self.set_cell_background(cell, "2F5496")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(6)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(12)
                    run.font.bold = True

        # Data rows
        row_colors = ["FFFFFF", "F2F2F2"]
        for row_idx, label in enumerate(labels, start=1):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = label
            for run in row_cells[0].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)

            bg_color = row_colors[row_idx % 2]
            for cell in row_cells:
                self.set_cell_background(cell, bg_color)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(6)

            row_data = data[row_idx - 1]
            for col_idx, value in enumerate(row_data, start=1):
                row_cells[col_idx].text = str(value)

    def generate(self):
        self.doc = Document()
        # Title
        title = self.doc.add_heading(self.corridor, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Traffic section
        self.add_formatted_section(self.section[0], 9)
        self.add_formatted_para("Section Length: \t", self.parameters['section_length'], "km")
        self.add_formatted_para("Train Average Speed: \t", self.parameters['average_speed'], "km/h")
        self.add_formatted_para("Train Composition: \t", self.train_composition)
        self.add_formatted_para("Train Carrying Capacity (AW3 Load): \t", self.capacity)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(18)

        traffic_data = [
            [str(self.data['DailyRidership'][yr]) for yr in self.years],
            [str(self.data['PHPDT'][yr]) for yr in self.years],
            [str(round(self.yearly_headways[yr], 1)) for yr in self.years],
            [str(self.yearly_trains[yr]) for yr in self.years]
        ]
        self.create_formatted_table(self.years, self.tfc_labels, traffic_data)

        # Plot
        self.doc.add_heading("Transit Data Plot", level=1)
        self.doc.add_picture(self.image_filename, width=Inches(6))
        self.doc.add_paragraph("The above plot shows the normalized transit data over the specified years.")

        # Power section
        self.add_formatted_section(self.section[1], 9)
        self.add_formatted_para("Specific Energy Consumption: \t", self.tpower['SEC'], "KWH/GTKM")
        self.add_formatted_para("Regeneration Percentage: \t", self.tpower['Regen'], "%")
        self.add_formatted_para("Losses in Traction Power: \t", self.tpower['TrLoss'], "%")
        self.add_formatted_para("Traction Power Factor: \t", self.tpower['TrPF'])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(18)

        self.add_formatted_para("Number of Elevated Station: \t", self.apower['ElStnNos'])
        self.add_formatted_para("Number of Under Ground Station: \t", self.apower['UGStnNos'])
        self.add_formatted_para("Number of Depots: \t", self.apower['DpNos'])
        self.add_formatted_para("Losses in Auxiliary Power: \t", self.apower['AuxLoss'], "%")
        self.add_formatted_para("Auxiliary Power Factor: \t", self.apower['AuxPF'])

        energy_data = [
            [str(round(self.trc_energy[0][yr], 2)) for yr in self.years],
            [str(round(self.trc_energy[1][yr], 2)) for yr in self.years],
            [str(round(self.aux_energy[0][yr], 2)) for yr in self.years],
            [str(round(self.aux_energy[1][yr], 2)) for yr in self.years],
            [str(round(self.total_energy[0][yr], 2)) for yr in self.years],
            [str(round(self.total_energy[1][yr], 2)) for yr in self.years]
        ]
        self.create_formatted_table(self.years, self.pw_labels, energy_data)

        # Save
        self.doc.save(self.doc_path)
        return self.doc_path
